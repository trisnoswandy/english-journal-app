import streamlit as st
from google import genai
from google.genai import types
from PIL import Image
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from io import BytesIO
from datetime import datetime
import re

try:
    import yfinance as yf
    YFINANCE_AVAILABLE = True
except ImportError:
    YFINANCE_AVAILABLE = False

st.set_page_config(page_title="Trisno's Language Lab", page_icon="🌴", layout="wide")

# ----------------------------------------------------------------------------
# STYLING
# ----------------------------------------------------------------------------
custom_css = """
<style>
    .stApp {
        background-color: #050b14;
        background-image:
            radial-gradient(circle at 50% 10%, rgba(0, 255, 170, 0.08) 0%, transparent 60%),
            linear-gradient(rgba(0, 255, 170, 0.03) 1px, transparent 1px),
            linear-gradient(90deg, rgba(0, 255, 170, 0.03) 1px, transparent 1px);
        background-size: 100% 100%, 40px 40px, 40px 40px;
        color: #f1f5f9;
    }
    .footer {
        position: fixed; left: 0; bottom: 0; width: 100%;
        background-color: rgba(5, 11, 20, 0.95); color: #94a3b8;
        text-align: center; padding: 10px; font-size: 12px;
        border-top: 1px solid #1e293b; z-index: 100; backdrop-filter: blur(5px);
    }
    h1, h2, h3 { color: #00ffaa !important; font-family: 'Courier New', Courier, monospace; }
    .stButton>button {
        background-color: rgba(0, 255, 170, 0.05) !important;
        color: #00ffaa !important; border: 1px solid #00ffaa !important;
        border-radius: 8px; font-weight: bold;
    }
    .stButton>button:hover {
        background-color: #00ffaa !important; color: #050b14 !important;
    }
    .metric-card {
        background-color: rgba(0, 255, 170, 0.05);
        border: 1px solid #1e293b;
        border-radius: 10px;
        padding: 14px;
    }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# ----------------------------------------------------------------------------
# SESSION STATE
# ----------------------------------------------------------------------------
defaults = {
    "eval_result": None,
    "answer_feedback": None,
    "history": [],          # list of dicts: {timestamp, language, score, input_preview, eval_result}
    "targets": [],          # list of dicts: {text, done}
    "api_key_valid": None,
}
for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

LANGUAGES = [
    "English", "Spanish", "French", "German", "Japanese",
    "Korean", "Mandarin Chinese", "Arabic", "Italian", "Portuguese",
]

MODEL_OPTIONS = {
    "Gemini 3.6 Flash (terbaru, seimbang)": "gemini-3.6-flash",
    "Gemini 3.5 Flash-Lite (paling hemat/cepat)": "gemini-3.5-flash-lite",
    "Gemini 3.1 Pro (paling teliti, lebih lambat/mahal)": "gemini-3.1-pro",
}

# Watchlist emiten sawit yang tercatat di Bursa Efek Indonesia (kode + suffix .JK
# yang dipakai yfinance untuk mengambil data dari Yahoo Finance).
PALM_TICKERS = {
    "Astra Agro Lestari (AALI)": "AALI.JK",
    "PP London Sumatra Indonesia (LSIP)": "LSIP.JK",
    "Salim Ivomas Pratama (SIMP)": "SIMP.JK",
    "Tunas Baru Lampung (TBLA)": "TBLA.JK",
    "Dharma Satya Nusantara (DSNG)": "DSNG.JK",
    "Sampoerna Agro (SGRO)": "SGRO.JK",
    "Sawit Sumbermas Sarana (SSMS)": "SSMS.JK",
    "Eagle High Plantations (BWPT)": "BWPT.JK",
}

# ----------------------------------------------------------------------------
# HELPERS
# ----------------------------------------------------------------------------

def extract_score(text: str):
    """Best-effort extraction of a numeric score/percentage from the AI response."""
    if not text:
        return None
    match = re.search(r"(\d{1,3})\s*%|\bskor[:\s]+(\d{1,3})", text, flags=re.IGNORECASE)
    if match:
        val = match.group(1) or match.group(2)
        try:
            n = int(val)
            return n if 0 <= n <= 100 else None
        except ValueError:
            return None
    return None


def build_system_prompt(language: str) -> str:
    return f"""
Bertindaklah sebagai Mentor Bahasa Pribadi untuk pembelajar yang sedang belajar {language}.
Evaluasi draf jurnal pengguna dan berikan umpan balik yang terstruktur persis mengikuti
12 Poin Evaluasi Jurnal Harian berikut. Seluruh contoh, koreksi, dan latihan HARUS dalam
bahasa target ({language}), dengan penjelasan tambahan berbahasa Indonesia agar mudah dipahami.

1. **Transkripsi Teks & Versi Alami (Natural)**
   Tampilkan draf asli pengguna dan versi perbaikan berbahasa {language} yang alami, dalam paragraf rapi.

2. **Kamus Kata-Kata (Siap Salin)**
   Daftar kosakata baru dari draf pengguna (format rapat, tanpa tabel) beserta arti dan contoh kalimat,
   siap ditempel ke Notion.

3. **Bedah Struktur & Pilihan Kata (Word-by-Word Analysis)**
   Analisis mendalam kesalahan tata bahasa, ejaan, tata kata, dilengkapi panduan cara baca fonetiknya.

4. **Before vs After Transformation**
   Perbandingan kalimat kurang tepat pada draf asli dengan versi perbaikannya beserta pelajaran utamanya.

5. **Top 3 Vocabulary Focus**
   Tiga kosakata atau idiom paling penting dari jurnal hari itu yang wajib dihafalkan.

6. **Mastered Verb/Conjugation Tracker**
   Pelacakan penggunaan konjugasi/kata kerja untuk melihat mana yang sudah dikuasai dan mana yang perlu diperbaiki.

7. **Skor & Persentase Ketepatan Mandiri**
   Cantumkan skor akurasi tulisan secara eksplisit dalam format "Skor: XX%" menggunakan formula:
   Skor = (Kata Benar / Total Kata) x 100%

8. **Pronunciation Challenge**
   Kalimat terpilih dari jurnal untuk dilatih secara lisan, lengkap dengan panduan pengucapan alami.

9. **Native Phrase of the Day**
   Satu frasa/idiom khas penutur asli {language} beserta contoh penggunaannya dalam kalimat.

10. **Evaluasi Jurnal Harian**
    Umpan balik ringkas, catatan perkembangan, dan apresiasi atas alur cerita jurnal pengguna.

11. **Daily Micro-Question**
    Satu pertanyaan singkat berbahasa {language} yang relevan dengan topik jurnal, untuk melatih respons cepat.

12. **Sesi Belajar Singkat (Materi & Latihan Soal)**
    Pembahasan materi dasar secara berurutan (6 Poin Fondasi) dengan siklus 3 hari pengulangan per materi,
    diakhiri 3 latihan soal singkat (Soal 1, Soal 2, Soal 3).
"""


def call_gemini(api_key: str, model: str, contents):
    """Call Gemini with a simple fallback to a lighter model if the primary fails."""
    client = genai.Client(api_key=api_key)
    try:
        return client.models.generate_content(model=model, contents=contents)
    except Exception as primary_error:
        fallback_model = "gemini-2.5-flash"
        if model == fallback_model:
            raise primary_error
        try:
            return client.models.generate_content(model=fallback_model, contents=contents)
        except Exception:
            raise primary_error


@st.cache_data(ttl=900, show_spinner=False)
def get_price_history(ticker: str, period: str = "6mo") -> pd.DataFrame:
    """Ambil data harga historis dari Yahoo Finance via yfinance. Cache 15 menit."""
    if not YFINANCE_AVAILABLE:
        return pd.DataFrame()
    try:
        data = yf.Ticker(ticker).history(period=period)
        return data
    except Exception:
        return pd.DataFrame()


@st.cache_data(ttl=900, show_spinner=False)
def get_fundamentals(ticker: str) -> dict:
    """Ambil ringkasan fundamental (P/E, market cap, dsb) dari Yahoo Finance."""
    if not YFINANCE_AVAILABLE:
        return {}
    try:
        info = yf.Ticker(ticker).info
        return {
            "Nama": info.get("longName", ticker),
            "Harga Terakhir": info.get("currentPrice") or info.get("regularMarketPrice"),
            "Mata Uang": info.get("currency", "-"),
            "P/E Ratio (TTM)": info.get("trailingPE"),
            "Market Cap": info.get("marketCap"),
            "Dividend Yield (%)": (
                round(info.get("dividendYield") * 100, 2)
                if info.get("dividendYield") else None
            ),
            "52-Week High": info.get("fiftyTwoWeekHigh"),
            "52-Week Low": info.get("fiftyTwoWeekLow"),
            "Sektor": info.get("sector", "-"),
        }
    except Exception:
        return {}


def search_grounded(api_key: str, model: str, query: str):
    """Panggil Gemini dengan Google Search grounding agar jawaban berbasis
    berita/data real-time, lengkap dengan sumbernya (bukan hasil karangan model)."""
    client = genai.Client(api_key=api_key)
    grounding_tool = types.Tool(google_search=types.GoogleSearch())
    config = types.GenerateContentConfig(tools=[grounding_tool])
    response = client.models.generate_content(model=model, contents=query, config=config)

    sources = []
    try:
        for candidate in response.candidates:
            gm = getattr(candidate, "grounding_metadata", None)
            if gm and gm.grounding_chunks:
                for chunk in gm.grounding_chunks:
                    if getattr(chunk, "web", None):
                        sources.append((chunk.web.title, chunk.web.uri))
    except Exception:
        pass
    return response.text, sources


def build_docx(title: str, body_text: str) -> BytesIO:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Dibuat pada: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph("")
    for line in body_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped.replace("### ", ""), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped.replace("## ", ""), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped.replace("# ", ""), level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ----------------------------------------------------------------------------
# SIDEBAR
# ----------------------------------------------------------------------------
with st.sidebar:
    st.markdown("### 🌴 Trisno's Language Lab")
    st.caption("Intelligence & Language Learning Platform")
    st.markdown("---")

    api_key = st.text_input("Gemini API Key", type="password", placeholder="masukkan API key Anda")
    st.caption("🔒 Key hanya dipakai di sesi ini dan tidak disimpan di server.")

    target_language = st.selectbox("🌐 Bahasa yang dipelajari:", LANGUAGES, index=0)

    model_label = st.selectbox("🤖 Model AI:", list(MODEL_OPTIONS.keys()), index=0)
    MODEL_NAME = MODEL_OPTIONS[model_label]

    st.markdown("---")
    menu = st.radio(
        "Navigasi:",
        [
            "📝 12-Point Journal Evaluator",
            "📈 Progress & Target Planner",
            "🎬 AI Video Prompt Gen",
            "🗂️ Riwayat Evaluasi",
            "🌴 Riset Perkebunan Sawit",
            "💹 Pasar Modal & Berita Saham",
        ],
    )
    st.markdown("---")

    total_sessions = len(st.session_state.history)
    st.markdown(f"**Sesi selesai:** {total_sessions}")
    st.markdown("👤 **Trisno Swandy Simanullang**")

# ----------------------------------------------------------------------------
# HEADER
# ----------------------------------------------------------------------------
st.title("🌴 Trisno's Language Lab")
st.markdown(
    "Platform Cerdas Pengembang Bahasa Asing, Riset Perkebunan Sawit, "
    "& Analisis Pasar Modal / Saham."
)

if not api_key:
    st.warning("⚠️ Masukkan Gemini API Key Anda di sidebar untuk mengaktifkan fitur AI.")

# ----------------------------------------------------------------------------
# MENU 1: JOURNAL EVALUATOR
# ----------------------------------------------------------------------------
if menu == "📝 12-Point Journal Evaluator":
    st.header(f"📝 12-Point Journal Evaluator — {target_language}")

    col1, col2 = st.columns([1, 1])
    with col1:
        journal_input = st.text_area(
            f"Input Teks Jurnal / Catatan Riset ({target_language}/ID):", height=180
        )
    with col2:
        uploaded_image = st.file_uploader(
            "Upload Foto / Screenshot Jurnal (PNG, JPG, JPEG):", type=["png", "jpg", "jpeg"]
        )
        if uploaded_image:
            image_preview = Image.open(uploaded_image)
            st.image(image_preview, caption="Preview Foto SS Jurnal", use_container_width=True)

    if st.button("🚀 INITIATE EVALUATION"):
        if not api_key:
            st.error("API Key belum dimasukkan di sidebar!")
        elif not journal_input and not uploaded_image:
            st.warning("Mohon masukkan teks jurnal atau upload foto screenshot terlebih dahulu.")
        else:
            with st.spinner(f"Menganalisis 12 Poin Evaluasi Jurnal ({target_language})..."):
                try:
                    system_prompt = build_system_prompt(target_language)
                    contents_payload = [system_prompt]
                    if journal_input:
                        contents_payload.append(f"\nTeks Jurnal Input:\n{journal_input}")
                    if uploaded_image:
                        contents_payload.append(Image.open(uploaded_image))

                    response = call_gemini(api_key, MODEL_NAME, contents_payload)

                    st.session_state.eval_result = response.text
                    st.session_state.answer_feedback = None

                    st.session_state.history.append({
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "language": target_language,
                        "score": extract_score(response.text),
                        "input_preview": (journal_input[:80] + "...") if journal_input else "[foto/screenshot]",
                        "eval_result": response.text,
                    })
                except Exception as e:
                    st.error(
                        "Terjadi kesalahan saat menghubungi Gemini API. "
                        "Periksa kembali API Key dan koneksi Anda.\n\n"
                        f"Detail teknis: {str(e)}"
                    )

    if st.session_state.eval_result:
        st.success("Evaluasi 12 Poin Selesai!")
        st.markdown(st.session_state.eval_result)

        docx_buffer = build_docx(
            f"Evaluasi Jurnal - {target_language} - {datetime.now().strftime('%Y-%m-%d')}",
            st.session_state.eval_result,
        )
        st.download_button(
            "📄 Unduh Hasil Evaluasi (.docx)",
            data=docx_buffer,
            file_name=f"evaluasi_jurnal_{target_language.lower()}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        st.markdown("---")
        st.subheader("✍️ Lembar Jawaban Latihan Soal & Daily Micro-Question")
        st.caption("Jawablah pertanyaan dari poin 11 dan 12 di atas untuk diperiksa langsung oleh AI.")

        user_answers = st.text_area(
            "Tuliskan jawaban Anda di sini (misal: Jawaban Micro-Question, Jawaban Soal 1, 2, dan 3):",
            height=150,
            placeholder="Contoh:\nDaily Micro-Question: I usually study in the evening...\nSoal 1: B\nSoal 2: went\nSoal 3: because",
        )

        if st.button("✅ Periksa Jawaban Saya"):
            if not api_key:
                st.error("API Key belum dimasukkan!")
            elif not user_answers:
                st.warning("Silakan tuliskan jawaban Anda terlebih dahulu.")
            else:
                with st.spinner("Mengevaluasi jawaban Anda..."):
                    try:
                        check_prompt = f"""
                        Berikut adalah konteks soal dan hasil evaluasi sebelumnya (bahasa target: {target_language}):
                        {st.session_state.eval_result}

                        Berikut adalah jawaban dari pengguna:
                        "{user_answers}"

                        Tugasmu:
                        1. Berikan koreksi dan penilaian apakah jawaban pengguna untuk Daily Micro-Question
                           dan 3 Latihan Soal sudah benar.
                        2. Jika ada yang salah, jelaskan letak kesalahannya dan berikan jawaban yang benar
                           beserta penjelasannya dalam bahasa Indonesia & {target_language}.
                        3. Cantumkan nilai/skor akhir untuk sesi latihan ini secara eksplisit dalam format "Skor: XX%".
                        """
                        feedback_res = call_gemini(api_key, MODEL_NAME, check_prompt)
                        st.session_state.answer_feedback = feedback_res.text
                    except Exception as e:
                        st.error(f"Terjadi kesalahan saat memeriksa jawaban: {str(e)}")

        if st.session_state.answer_feedback:
            st.markdown("### 📊 Hasil Penilaian Jawaban Anda")
            st.info(st.session_state.answer_feedback)

# ----------------------------------------------------------------------------
# MENU 2: PROGRESS & TARGET PLANNER
# ----------------------------------------------------------------------------
elif menu == "📈 Progress & Target Planner":
    st.header("📈 Progress & Target Planner")

    tab_targets, tab_progress = st.tabs(["🎯 Target Belajar", "📊 Grafik Progres"])

    with tab_targets:
        st.caption("Tambahkan target belajar Anda dan centang setelah tercapai.")
        with st.form("add_target_form", clear_on_submit=True):
            new_target = st.text_input("Target baru:", placeholder="Contoh: Kuasai 20 kata kerja tidak beraturan")
            submitted = st.form_submit_button("➕ Tambah Target")
            if submitted and new_target:
                st.session_state.targets.append({"text": new_target, "done": False})

        if not st.session_state.targets:
            st.info("Belum ada target. Tambahkan target pertama Anda di atas.")
        else:
            done_count = sum(1 for t in st.session_state.targets if t["done"])
            st.progress(done_count / len(st.session_state.targets))
            st.caption(f"{done_count} dari {len(st.session_state.targets)} target tercapai.")

            for i, t in enumerate(st.session_state.targets):
                col_check, col_del = st.columns([8, 1])
                with col_check:
                    t["done"] = st.checkbox(t["text"], value=t["done"], key=f"target_{i}")
                with col_del:
                    if st.button("🗑️", key=f"del_{i}"):
                        st.session_state.targets.pop(i)
                        st.rerun()

    with tab_progress:
        if not st.session_state.history:
            st.info("Belum ada data evaluasi. Selesaikan sesi di '12-Point Journal Evaluator' dulu.")
        else:
            df = pd.DataFrame(st.session_state.history)
            df["session_no"] = range(1, len(df) + 1)

            col_a, col_b, col_c = st.columns(3)
            with col_a:
                st.markdown(f"<div class='metric-card'><b>Total Sesi</b><br>{len(df)}</div>", unsafe_allow_html=True)
            with col_b:
                avg_score = df["score"].dropna().mean()
                avg_display = f"{avg_score:.0f}%" if pd.notna(avg_score) else "N/A"
                st.markdown(f"<div class='metric-card'><b>Rata-rata Skor</b><br>{avg_display}</div>", unsafe_allow_html=True)
            with col_c:
                st.markdown(f"<div class='metric-card'><b>Bahasa Aktif</b><br>{target_language}</div>", unsafe_allow_html=True)

            if df["score"].notna().any():
                fig = px.line(
                    df, x="session_no", y="score", markers=True, color="language",
                    labels={"session_no": "Sesi ke-", "score": "Skor (%)"},
                    title="Perkembangan Skor per Sesi",
                )
                fig.update_layout(
                    plot_bgcolor="#050b14", paper_bgcolor="#050b14", font_color="#f1f5f9"
                )
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.caption("Skor belum terdeteksi otomatis dari hasil evaluasi sebelumnya.")

            st.dataframe(
                df[["timestamp", "language", "score", "input_preview"]],
                use_container_width=True,
                hide_index=True,
            )

    st.caption(
        "ℹ️ Catatan: data progres tersimpan hanya selama sesi browser ini berjalan "
        "(khas aplikasi Streamlit tanpa database). Unduh hasil evaluasi sebagai .docx "
        "jika ingin menyimpannya secara permanen."
    )

# ----------------------------------------------------------------------------
# MENU 3: AI VIDEO PROMPT GENERATOR
# ----------------------------------------------------------------------------
elif menu == "🎬 AI Video Prompt Gen":
    st.header("🎬 AI Video Prompt Generator")

    prompt_input = st.text_area("Deskripsi Visual / Ide Kreatif:", height=120)

    col1, col2, col3 = st.columns(3)
    with col1:
        style = st.selectbox("Gaya visual:", ["Cinematic", "Documentary", "Anime", "Hyperrealistic", "Stop-motion"])
    with col2:
        duration = st.selectbox("Durasi:", ["5 detik", "10 detik", "15 detik", "30 detik"])
    with col3:
        camera = st.selectbox("Gerakan kamera:", ["Static", "Slow pan", "Dolly zoom", "Drone aerial", "Handheld"])

    if st.button("✨ Generate Prompt"):
        if not api_key:
            st.error("API Key belum dimasukkan!")
        elif not prompt_input:
            st.warning("Masukkan deskripsi visual terlebih dahulu.")
        else:
            with st.spinner("Menyusun prompt video..."):
                try:
                    full_instruction = (
                        f"Ubah ide berikut menjadi prompt video AI yang sinematik dan detail, "
                        f"dalam bahasa Inggris, siap pakai untuk model text-to-video:\n\n"
                        f"Ide: {prompt_input}\n"
                        f"Gaya visual: {style}\n"
                        f"Durasi: {duration}\n"
                        f"Gerakan kamera: {camera}\n\n"
                        f"Sertakan detail pencahayaan, suasana, dan komposisi shot."
                    )
                    response = call_gemini(api_key, MODEL_NAME, full_instruction)
                    st.success("Prompt Berhasil Dibuat!")
                    st.code(response.text, language="markdown")
                except Exception as e:
                    st.error(f"Terjadi kesalahan: {str(e)}")

# ----------------------------------------------------------------------------
# MENU 4: HISTORY
# ----------------------------------------------------------------------------
elif menu == "🗂️ Riwayat Evaluasi":
    st.header("🗂️ Riwayat Evaluasi")

    if not st.session_state.history:
        st.info("Belum ada riwayat evaluasi pada sesi ini.")
    else:
        for i, item in enumerate(reversed(st.session_state.history)):
            score_label = f" — Skor: {item['score']}%" if item.get("score") is not None else ""
            with st.expander(f"{item['timestamp']} | {item['language']}{score_label} | {item['input_preview']}"):
                st.markdown(item["eval_result"])
                docx_buffer = build_docx(
                    f"Evaluasi Jurnal - {item['language']} - {item['timestamp']}",
                    item["eval_result"],
                )
                st.download_button(
                    "📄 Unduh sebagai .docx",
                    data=docx_buffer,
                    file_name=f"evaluasi_{item['timestamp'].replace(' ', '_').replace(':', '')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"hist_dl_{i}",
                )

# ----------------------------------------------------------------------------
# MENU 5: RISET PERKEBUNAN SAWIT
# ----------------------------------------------------------------------------
elif menu == "🌴 Riset Perkebunan Sawit":
    st.header("🌴 Riset Perkebunan Sawit")

    tab_diagnosa, tab_emiten, tab_berita = st.tabs(
        ["🩺 Diagnosa Tanaman (AI)", "📊 Data Emiten Sawit", "📰 Berita & Harga CPO"]
    )

    with tab_diagnosa:
        st.caption(
            "Upload foto daun, buah, atau kondisi tanaman sawit untuk mendapat "
            "analisis awal dari AI mengenai kemungkinan penyakit/hama."
        )
        leaf_image = st.file_uploader(
            "Upload foto tanaman sawit:", type=["png", "jpg", "jpeg"], key="palm_leaf"
        )
        extra_notes = st.text_area(
            "Catatan tambahan (opsional): usia tanaman, gejala, lokasi kebun, dll.",
            height=100,
        )
        if leaf_image:
            st.image(Image.open(leaf_image), caption="Preview foto", use_container_width=True)

        if st.button("🔍 Analisis Kondisi Tanaman"):
            if not api_key:
                st.error("API Key belum dimasukkan!")
            elif not leaf_image:
                st.warning("Upload foto tanaman terlebih dahulu.")
            else:
                with st.spinner("Menganalisis kondisi tanaman..."):
                    try:
                        diag_prompt = (
                            "Bertindaklah sebagai agronomis spesialis kelapa sawit. "
                            "Amati foto berikut dan berikan: (1) kemungkinan kondisi/penyakit/hama "
                            "yang teramati beserta tingkat keyakinan, (2) gejala visual yang mendukung "
                            "dugaan tersebut, (3) rekomendasi langkah penanganan awal, "
                            "(4) kapan sebaiknya berkonsultasi dengan agronomis atau PPL setempat. "
                            "Jawab dalam Bahasa Indonesia, dan tegaskan bahwa ini adalah analisis awal "
                            "berbasis AI, bukan pengganti pemeriksaan lapangan oleh ahli."
                        )
                        payload = [diag_prompt, Image.open(leaf_image)]
                        if extra_notes:
                            payload.append(f"Catatan tambahan dari pengguna: {extra_notes}")
                        response = call_gemini(api_key, MODEL_NAME, payload)
                        st.success("Analisis selesai!")
                        st.markdown(response.text)
                        st.caption(
                            "⚠️ Ini adalah bantuan awal berbasis AI dan bisa keliru. "
                            "Untuk keputusan penanganan/penyemprotan, konfirmasikan ke agronomis "
                            "atau petugas penyuluh lapangan (PPL) sebelum bertindak."
                        )
                    except Exception as e:
                        st.error(f"Terjadi kesalahan: {str(e)}")

    with tab_emiten:
        if not YFINANCE_AVAILABLE:
            st.error(
                "Modul `yfinance` belum terpasang di environment ini. "
                "Tambahkan `yfinance` ke requirements.txt lalu redeploy aplikasi."
            )
        else:
            selected_company = st.selectbox("Pilih emiten sawit (IDX):", list(PALM_TICKERS.keys()))
            ticker = PALM_TICKERS[selected_company]
            period = st.select_slider(
                "Periode grafik:", options=["1mo", "3mo", "6mo", "1y", "2y", "5y"], value="6mo"
            )

            with st.spinner(f"Mengambil data {ticker} dari Yahoo Finance..."):
                hist = get_price_history(ticker, period)
                fundamentals = get_fundamentals(ticker)

            if hist.empty:
                st.warning(
                    f"Data untuk {ticker} tidak ditemukan atau gagal diambil. "
                    "Coba lagi beberapa saat lagi, atau periksa apakah ticker masih aktif."
                )
            else:
                fig = go.Figure(
                    data=[go.Candlestick(
                        x=hist.index, open=hist["Open"], high=hist["High"],
                        low=hist["Low"], close=hist["Close"],
                    )]
                )
                fig.update_layout(
                    title=f"Pergerakan Harga {selected_company} ({ticker})",
                    plot_bgcolor="#050b14", paper_bgcolor="#050b14", font_color="#f1f5f9",
                    xaxis_rangeslider_visible=False,
                )
                st.plotly_chart(fig, use_container_width=True)

                if fundamentals:
                    st.subheader("📋 Ringkasan Fundamental")
                    fcol1, fcol2, fcol3 = st.columns(3)
                    items = list(fundamentals.items())
                    for i, (label, value) in enumerate(items):
                        target_col = [fcol1, fcol2, fcol3][i % 3]
                        display_val = "-" if value is None else value
                        target_col.markdown(
                            f"<div class='metric-card'><b>{label}</b><br>{display_val}</div>",
                            unsafe_allow_html=True,
                        )
                st.caption(
                    "Sumber data: Yahoo Finance (via `yfinance`). Data bisa delay 15-20 menit "
                    "dan bukan rekomendasi investasi."
                )

    with tab_berita:
        st.caption(
            "Menggunakan Gemini dengan Google Search grounding agar jawaban berbasis "
            "berita/data real-time, lengkap dengan sumber aslinya."
        )
        default_query = "Perkembangan harga CPO (Crude Palm Oil) dan industri sawit Indonesia minggu ini"
        news_query = st.text_input("Topik pencarian:", value=default_query)

        if st.button("📰 Cari & Rangkum Berita"):
            if not api_key:
                st.error("API Key belum dimasukkan!")
            else:
                with st.spinner("Mencari dan merangkum berita terbaru..."):
                    try:
                        text, sources = search_grounded(api_key, MODEL_NAME, news_query)
                        st.markdown(text)
                        if sources:
                            st.markdown("**Sumber:**")
                            for title, uri in sources:
                                st.markdown(f"- [{title or uri}]({uri})")
                        else:
                            st.caption("Tidak ada metadata sumber yang dikembalikan untuk kueri ini.")
                    except Exception as e:
                        st.error(f"Terjadi kesalahan saat mencari berita: {str(e)}")

# ----------------------------------------------------------------------------
# MENU 6: PASAR MODAL & BERITA SAHAM
# ----------------------------------------------------------------------------
elif menu == "💹 Pasar Modal & Berita Saham":
    st.header("💹 Pasar Modal & Berita Saham")

    tab_watch, tab_detail, tab_news = st.tabs(
        ["👀 Watchlist", "📈 Detail Saham", "📰 Berita Pasar (AI)"]
    )

    with tab_watch:
        if not YFINANCE_AVAILABLE:
            st.error(
                "Modul `yfinance` belum terpasang di environment ini. "
                "Tambahkan `yfinance` ke requirements.txt lalu redeploy aplikasi."
            )
        else:
            st.caption(
                "Bandingkan performa beberapa saham sekaligus (perubahan harga dalam %, "
                "dinormalisasi dari awal periode)."
            )
            default_tickers = "AALI.JK, BBCA.JK, TLKM.JK"
            tickers_input = st.text_input(
                "Kode saham (pisahkan dengan koma, gunakan .JK untuk saham IDX):",
                value=default_tickers,
            )
            period = st.select_slider(
                "Periode:", options=["1mo", "3mo", "6mo", "1y", "2y"], value="6mo", key="watch_period"
            )

            if st.button("📊 Tampilkan Perbandingan"):
                tickers = [t.strip().upper() for t in tickers_input.split(",") if t.strip()]
                if not tickers:
                    st.warning("Masukkan minimal satu kode saham.")
                else:
                    fig = go.Figure()
                    any_data = False
                    for tk in tickers:
                        hist = get_price_history(tk, period)
                        if hist.empty:
                            st.warning(f"Data untuk {tk} tidak ditemukan.")
                            continue
                        normalized = (hist["Close"] / hist["Close"].iloc[0] - 1) * 100
                        fig.add_trace(go.Scatter(x=hist.index, y=normalized, mode="lines", name=tk))
                        any_data = True
                    if any_data:
                        fig.update_layout(
                            title="Perbandingan Kinerja Saham (% perubahan)",
                            plot_bgcolor="#050b14", paper_bgcolor="#050b14", font_color="#f1f5f9",
                            yaxis_title="Perubahan (%)",
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    st.caption("Sumber data: Yahoo Finance. Bukan rekomendasi investasi.")

    with tab_detail:
        if not YFINANCE_AVAILABLE:
            st.error("Modul `yfinance` belum terpasang. Tambahkan ke requirements.txt.")
        else:
            ticker_detail = st.text_input(
                "Kode saham (contoh: BBCA.JK, AALI.JK, AAPL untuk saham AS):", value="AALI.JK"
            )
            period_detail = st.select_slider(
                "Periode grafik:", options=["1mo", "3mo", "6mo", "1y", "2y", "5y"],
                value="1y", key="detail_period",
            )
            if st.button("🔎 Tampilkan Detail"):
                with st.spinner(f"Mengambil data {ticker_detail}..."):
                    hist = get_price_history(ticker_detail, period_detail)
                    fundamentals = get_fundamentals(ticker_detail)

                if hist.empty:
                    st.warning(
                        f"Data untuk '{ticker_detail}' tidak ditemukan. "
                        "Pastikan kode saham benar (gunakan .JK untuk saham IDX)."
                    )
                else:
                    fig = go.Figure(
                        data=[go.Candlestick(
                            x=hist.index, open=hist["Open"], high=hist["High"],
                            low=hist["Low"], close=hist["Close"],
                        )]
                    )
                    fig.update_layout(
                        title=f"Pergerakan Harga {ticker_detail}",
                        plot_bgcolor="#050b14", paper_bgcolor="#050b14", font_color="#f1f5f9",
                        xaxis_rangeslider_visible=False,
                    )
                    st.plotly_chart(fig, use_container_width=True)

                    if fundamentals:
                        st.subheader("📋 Ringkasan Fundamental")
                        fcol1, fcol2, fcol3 = st.columns(3)
                        items = list(fundamentals.items())
                        for i, (label, value) in enumerate(items):
                            target_col = [fcol1, fcol2, fcol3][i % 3]
                            display_val = "-" if value is None else value
                            target_col.markdown(
                                f"<div class='metric-card'><b>{label}</b><br>{display_val}</div>",
                                unsafe_allow_html=True,
                            )
                    st.caption("Sumber data: Yahoo Finance (via `yfinance`). Bukan rekomendasi investasi.")

    with tab_news:
        st.caption(
            "Menggunakan Gemini dengan Google Search grounding untuk berita pasar modal "
            "real-time beserta sumbernya."
        )
        default_news_query = "Berita terbaru pasar saham Indonesia (IHSG) dan sentimen pasar hari ini"
        market_query = st.text_input("Topik pencarian:", value=default_news_query, key="market_news_query")

        if st.button("📰 Cari & Rangkum Berita Pasar"):
            if not api_key:
                st.error("API Key belum dimasukkan!")
            else:
                with st.spinner("Mencari dan merangkum berita pasar terbaru..."):
                    try:
                        text, sources = search_grounded(api_key, MODEL_NAME, market_query)
                        st.markdown(text)
                        if sources:
                            st.markdown("**Sumber:**")
                            for title, uri in sources:
                                st.markdown(f"- [{title or uri}]({uri})")
                        else:
                            st.caption("Tidak ada metadata sumber yang dikembalikan untuk kueri ini.")
                    except Exception as e:
                        st.error(f"Terjadi kesalahan saat mencari berita: {str(e)}")

        st.info(
            "⚠️ Informasi di atas dihasilkan AI berbasis hasil pencarian dan hanya untuk tujuan "
            "edukasi/riset. Ini bukan nasihat keuangan — selalu verifikasi ke sumber resmi "
            "(IDX, OJK, laporan keuangan emiten) sebelum mengambil keputusan investasi."
        )

# ----------------------------------------------------------------------------
# FOOTER
# ----------------------------------------------------------------------------
st.markdown(
    "<div class='footer'>Developed with ⚡ by <b>Trisno Swandy Simanullang</b></div>",
    unsafe_allow_html=True,
)
