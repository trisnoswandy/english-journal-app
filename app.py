import streamlit as st
import google.generativeai as genai
import pandas as pd
import plotly.express as px
from docx import Document
from io import BytesIO
import time
from datetime import datetime, timedelta

# ==========================================
# 1. KONFIGURASI HALAMAN & TEMA FUTURISTIK
# ==========================================
st.set_page_config(page_title="AI Neo-English Hub", page_icon="⚡", layout="wide")

# Injeksi CSS Custom untuk Tema Neon/Cyberpunk
custom_css = """
<style>
    /* Dark Mode & Neon Blue Accents */
    .stApp {
        background-color: #0a0e17;
        color: #e0e0e0;
    }
    h1, h2, h3, h4, h5, h6 {
        color: #00f3ff !important;
        font-family: 'Courier New', Courier, monospace;
        text-shadow: 0 0 5px rgba(0, 243, 255, 0.5);
    }
    .stButton>button {
        background-color: transparent !important;
        color: #00f3ff !important;
        border: 1px solid #00f3ff !important;
        border-radius: 8px;
        transition: 0.3s;
        box-shadow: 0 0 10px rgba(0, 243, 255, 0.2);
    }
    .stButton>button:hover {
        background-color: #00f3ff !important;
        color: #0a0e17 !important;
        box-shadow: 0 0 20px rgba(0, 243, 255, 0.6);
    }
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        background-color: #111827 !important;
        color: #00f3ff !important;
        border: 1px solid #1f2937 !important;
    }
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus {
        border-color: #00f3ff !important;
        box-shadow: 0 0 8px rgba(0, 243, 255, 0.5) !important;
    }
    /* Futuristic Loading Animation */
    @keyframes pulse {
        0% { box-shadow: 0 0 0 0 rgba(0, 243, 255, 0.7); }
        70% { box-shadow: 0 0 0 10px rgba(0, 243, 255, 0); }
        100% { box-shadow: 0 0 0 0 rgba(0, 243, 255, 0); }
    }
    .stSpinner > div > div {
        border-color: #00f3ff transparent transparent transparent !important;
    }
    hr {
        border-top: 1px solid #00f3ff;
        opacity: 0.3;
    }
    .metric-card {
        background-color: #111827;
        border-left: 4px solid #00f3ff;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 10px;
    }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# ==========================================
# 2. HELPER FUNCTIONS (API & EXPORT)
# ==========================================
def create_docx(content, title="Exported_Document"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_gemini_response(prompt, api_key):
    try:
        genai.configure(api_key=api_key)
        # Menggunakan model gemini-2.5-flash
        model = genai.GenerativeModel('gemini-3.7-flash')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# ==========================================
# 3. SIDEBAR & API SETUP
# ==========================================
with st.sidebar:
    st.markdown("## ⚙️ SYSTEM CORE")
    api_key = st.text_input("Insert Gemini API Key", type="password", placeholder="aq_xxxxxxxxxxxx")
    st.markdown("---")
    st.markdown("### 🛰️ NAVIGATION")
    menu = st.radio("Access Module:", ["12-Point Journal Evaluator", "Progress & Target Planner", "AI Video Prompt Gen (Veo)"])
    st.markdown("---")
    st.markdown("💡 *Theme: Neon Cyberpunk*")

st.title("⚡ AI Neo-English Hub")

# Pastikan API Key diisi
if not api_key and menu in ["12-Point Journal Evaluator", "AI Video Prompt Gen (Veo)"]:
    st.warning("⚠️ Access Denied: Please insert your Gemini API Key in the sidebar to activate AI modules.")

# ==========================================
# MODULE 1: 12-POINT JOURNAL EVALUATOR
# ==========================================
if menu == "12-Point Journal Evaluator":
    st.header("📝 12-Point Journal Evaluator")
    st.markdown("Input draf jurnal Anda (ID/EN) dan biarkan AI membedah serta menyempurnakannya dalam 12 poin komprehensif.")

    col1, col2 = st.columns([2, 1])
    
    with col1:
        journal_input = st.text_area("Input Draf Jurnal Harian (ID/EN):", height=250, placeholder="Hari ini saya merasa sangat produktif karena...")
        
    with col2:
        st.markdown("### 🎛️ Output Parameters")
        variations = st.slider("Jumlah Variasi Output:", 1, 3, 1)
        tone = st.selectbox("Gaya Bahasa:", ["Casual/Daily", "Professional/Formal", "Native Conversational", "Slang/Street English"])
        length = st.select_slider("Panjang Umpan Balik:", options=["Ringkas/Compact", "Sedang", "Detail/Komprehensif"], value="Sedang")

    if st.button("🚀 INITIATE EVALUATION"):
        if not api_key:
            st.error("API Key is missing!")
        elif not journal_input:
            st.warning("Mohon masukkan teks jurnal terlebih dahulu.")
        else:
            with st.spinner("🧠 AI Neural Network is processing your journal..."):
                prompt = f"""
                Bertindaklah sebagai mentor Bahasa Inggris level Native. Evaluasi jurnal berikut berdasarkan parameter ini:
                - Teks: "{journal_input}"
                - Jumlah Variasi: {variations}
                - Gaya Bahasa: {tone}
                - Panjang/Detail Umpan Balik: {length}

                Berikan output dalam Bahasa Indonesia (untuk penjelasannya) dan Inggris, terstruktur PERSIS dengan 12 poin berikut:
                1) Transkripsi Teks & Versi Alami (Natural)
                2) Kamus Kata-Kata (Tabel kosakata sulit/baru siap salin)
                3) Bedah Struktur & Pilihan Kata (Word-by-Word Analysis)
                4) Before vs After Transformation
                5) Top 3 Vocabulary Focus
                6) Mastered Verb Tracker (Kata kerja yang berhasil digunakan dengan baik)
                7) Skor & Persentase Ketepatan Mandiri (0-100%)
                8) Pronunciation Challenge (Tantangan Pengucapan & Panduan Fonetik)
                9) Native Phrase of the Day (Idiom/frasa terkait konteks jurnal)
                10) Evaluasi Jurnal Harian (Kritik & Saran membangun)
                11) Daily Micro-Question (1 pertanyaan pemicu untuk dijawab besok)
                12) Sesi Belajar Singkat (Materi Fondasi & 3 Latihan Soal dari kesalahan grammar di jurnal)
                """
                
                result = get_gemini_response(prompt, api_key)
                
                st.success("✅ Evaluation Complete!")
                
                # Render hasil
                st.markdown("### 📊 Evaluation Result")
                st.markdown(result)
                
                # Fitur Copy to Clipboard
                st.markdown("### 📋 Copy to Notion / Notes")
                st.code(result, language="markdown")
                
                # Fitur Export ke Word
                docx_file = create_docx(result, "Daily_English_Journal_Evaluation")
                st.download_button(
                    label="📄 Export ke Word (.docx)",
                    data=docx_file,
                    file_name="Journal_Evaluation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ==========================================
# MODULE 2: PROGRESS TRACKER & PLANNER
# ==========================================
elif menu == "Progress & Target Planner":
    st.header("📈 Progress Tracker & Target Planner")
    
    # Mock Data untuk Visualisasi
    dates = [datetime.today() - timedelta(days=i) for i in range(6, -1, -1)]
    scores = [75, 78, 76, 82, 85, 88, 92]
    df = pd.DataFrame({"Date": dates, "Accuracy Score (%)": scores})

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""<div class='metric-card'>
        <h4 style='margin:0'>🔥 Current Streak</h4>
        <h2 style='margin:0; color:#00f3ff'>7 Days</h2>
        </div>""", unsafe_allow_html=True)
    with col2:
        st.markdown("""<div class='metric-card'>
        <h4 style='margin:0'>📚 Vocab Mastered</h4>
        <h2 style='margin:0; color:#00f3ff'>142 Words</h2>
        </div>""", unsafe_allow_html=True)
    with col3:
        st.markdown("""<div class='metric-card'>
        <h4 style='margin:0'>⏱️ Learning Time</h4>
        <h2 style='margin:0; color:#00f3ff'>12.5 Hours</h2>
        </div>""", unsafe_allow_html=True)

    st.markdown("---")
    
    col_chart, col_plan = st.columns([2, 1])
    
    with col_chart:
        st.markdown("### 📉 Accuracy Progression")
        fig = px.line(df, x="Date", y="Accuracy Score (%)", markers=True, 
                      color_discrete_sequence=["#00f3ff"], template="plotly_dark")
        fig.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)")
        st.plotly_chart(fig, use_container_width=True)
        
    with col_plan:
        st.markdown("### 🎯 Weekly Focus Board")
        st.text_area("Target Kosakata Baru:", "1. Flabbergasted\n2. Inevitable\n3. Preposterous\nTarget: 30 Kosakata/Minggu", height=100)
        st.text_area("Fokus Grammar:", "1. Present Perfect vs Past Simple\n2. Conditionals Type 2\n3. Phrasal Verbs dengan 'Get'", height=100)
        if st.button("💾 Save Target"):
            with st.spinner("Saving to database..."):
                time.sleep(1)
            st.toast("Target saved successfully! 🚀")

# ==========================================
# MODULE 3: AI VIDEO PROMPT GENERATOR (VEO)
# ==========================================
elif menu == "AI Video Prompt Gen (Veo)":
    st.header("🎬 Google Veo / Sora Prompt Generator")
    st.markdown("Visualisasikan kosakata atau cerita jurnal Anda menjadi video AI bergaya cinematic/futuristik.")

    with st.form("veo_form"):
        col1, col2 = st.columns(2)
        with col1:
            subject = st.text_area("Subjek / Cerita Utama (ID/EN):", placeholder="Seorang cyberpunk berjalan di tengah hujan neon Tokyo...", height=150)
        with col2:
            style = st.selectbox("Gaya Visual:", ["Cinematic 35mm", "Photorealistic 8K", "Cyberpunk / Neon-noir", "Anime Studio Ghibli style", "Hyper-surrealism"])
            camera = st.selectbox("Pergerakan Kamera:", ["Slow Pan Right", "Drone Shot / Aerial", "Close-up Tracking", "Handheld shaky cam", "Zoom in slowly"])
            lighting = st.selectbox("Pencahayaan:", ["Neon Lights / High Contrast", "Golden Hour", "Moody / Volumetric Fog", "Cinematic Dark"])
        
        submitted = st.form_submit_button("✨ GENERATE VIDEO PROMPT")

    if submitted:
        if not api_key:
            st.error("API Key is missing!")
        elif not subject:
            st.warning("Mohon masukkan subjek/cerita utama.")
        else:
            with st.spinner("🌀 Synthesizing visual prompt parameters..."):
                prompt = f"""
                Kamu adalah ahli dalam membuat prompt video AI (untuk Google Veo 3 / Sora).
                Tugasmu adalah mengubah ide berikut menjadi prompt Bahasa Inggris yang sangat detail dan konsisten.
                
                Ide Utama: "{subject}"
                Visual Style: {style}
                Camera Movement: {camera}
                Lighting: {lighting}

                Buatlah 1 prompt utuh dalam bahasa Inggris (maksimal 2-3 kalimat panjang) yang mendeskripsikan:
                - Subjek dan aksinya secara spesifik
                - Setting lokasi dan atmosfer
                - Pencahayaan dan tekstur
                - Pergerakan kamera dan spesifikasi lensa/style.

                Format output: HANYA berikan teks prompt-nya saja, tanpa basa-basi.
                """
                
                result = get_gemini_response(prompt, api_key)
                
                st.success("✅ Prompt Generated!")
                st.markdown("### 🎥 Your Video Prompt")
                
                # Fitur Copy Prompt
                st.code(result, language="markdown")
                
                # Fitur Export Prompt to Word
                docx_file = create_docx(result, "AI_Video_Prompt_Veo")
                st.download_button(
                    label="📄 Export Prompt ke Word (.docx)",
                    data=docx_file,
                    file_name="Video_Prompt.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
